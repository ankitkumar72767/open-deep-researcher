from docx import Document

def create_docx(report):

    file_name = "research_report.docx"

    doc = Document()

    doc.add_heading(
        "Research Report",
        level=1
    )

    doc.add_paragraph(report)

    doc.save(file_name)

    return file_name